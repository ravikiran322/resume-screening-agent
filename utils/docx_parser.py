import docx
import os
import logging

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts all text from a DOCX file, including paragraphs and tables.
    
    Args:
        file_path (str): The absolute path to the DOCX file.
        
    Returns:
        str: The extracted text. Returns an empty string if extraction fails.
    """
    if not os.path.exists(file_path):
        logger.error(f"DOCX file not found: {file_path}")
        return ""
        
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
                
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in row_text:
                        row_text.append(cell_text)
                if row_text:
                    full_text.append(" | ".join(row_text))
                    
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
        return ""
