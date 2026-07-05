import os
import sys

# Ensure reportlab is available before generating PDFs
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except ImportError:
    print("reportlab is required to generate mock PDF files. Installing it...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "reportlab"])
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

try:
    import docx
except ImportError:
    print("python-docx is required to generate mock DOCX files. Installing it...")
    import subprocess
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

def create_txt_resume(filename, content):
    path = os.path.join("data", "resumes", filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"Created TXT resume: {path}")

def create_docx_resume(filename, name, email, phone, profile, experience, education, skills):
    path = os.path.join("data", "resumes", filename)
    doc = docx.Document()
    
    # Header
    doc.add_heading(name, 0)
    p = doc.add_paragraph()
    p.add_run(f"Email: {email} | Phone: {phone}\n")
    
    # Profile
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(profile)
    
    # Skills
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph(", ".join(skills))
    
    # Experience
    doc.add_heading("Work Experience", level=1)
    for job in experience:
        doc.add_heading(f"{job['title']} - {job['company']}", level=2)
        doc.add_paragraph(f"Duration: {job['duration']}")
        for resp in job['responsibilities']:
            doc.add_paragraph(resp, style='List Bullet')
            
    # Education
    doc.add_heading("Education", level=1)
    for edu in education:
        doc.add_paragraph(f"{edu['degree']} | {edu['institution']} ({edu['year']})")
        
    doc.save(path)
    print(f"Created DOCX resume: {path}")

def create_pdf_resume(filename, name, email, phone, profile, experience, education, skills):
    path = os.path.join("data", "resumes", filename)
    c = canvas.Canvas(path, pagesize=letter)
    width, height = letter
    
    # Draw simple styled resume
    y = height - 50
    
    # Title
    c.setFont("Helvetica-Bold", 24)
    c.drawString(50, y, name)
    y -= 20
    
    # Contact
    c.setFont("Helvetica", 10)
    c.drawString(50, y, f"Email: {email} | Phone: {phone}")
    y -= 15
    c.line(50, y, width - 50, y)
    y -= 25
    
    # Summary
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Professional Summary")
    y -= 15
    c.setFont("Helvetica", 10)
    # Split summary into lines if long
    summary_words = profile.split()
    lines = []
    current_line = []
    for word in summary_words:
        if len(" ".join(current_line + [word])) < 80:
            current_line.append(word)
        else:
            lines.append(" ".join(current_line))
            current_line = [word]
    if current_line:
        lines.append(" ".join(current_line))
        
    for line in lines:
        c.drawString(60, y, line)
        y -= 15
    y -= 10
    
    # Skills
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Technical Skills")
    y -= 15
    c.setFont("Helvetica", 10)
    skills_str = ", ".join(skills)
    c.drawString(60, y, skills_str)
    y -= 25
    
    # Experience
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Work Experience")
    y -= 15
    for job in experience:
        c.setFont("Helvetica-Bold", 11)
        c.drawString(60, y, f"{job['title']} - {job['company']}")
        c.setFont("Helvetica-Oblique", 10)
        c.drawString(width - 150, y, job['duration'])
        y -= 15
        c.setFont("Helvetica", 10)
        for resp in job['responsibilities'][:2]: # Take first two responsibilities to fit on page
            c.drawString(70, y, f"- {resp}")
            y -= 15
        y -= 5
    y -= 10
    
    # Education
    c.setFont("Helvetica-Bold", 14)
    c.drawString(50, y, "Education")
    y -= 15
    c.setFont("Helvetica", 10)
    for edu in education:
        c.drawString(60, y, f"{edu['degree']} - {edu['institution']} ({edu['year']})")
        y -= 15
        
    c.save()
    print(f"Created PDF resume: {path}")

def main():
    os.makedirs(os.path.join("data", "resumes"), exist_ok=True)
    
    # 1. Jane Smith (Data Scientist - High Match) - PDF
    create_pdf_resume(
        "Jane_Smith_Data_Scientist.pdf",
        "Jane Smith",
        "jane.smith@email.com",
        "+1-555-0101",
        "Experienced Data Scientist with a passion for extracting actionable insights from complex datasets. Proven track record of designing and deploying predictive machine learning models in cloud environments.",
        [
            {
                "title": "Senior Data Scientist",
                "company": "DataTech Solutions",
                "duration": "2021 - Present",
                "responsibilities": [
                    "Design and implement end-to-end Machine Learning pipelines using Python, Scikit-Learn, and PyTorch.",
                    "Deploy predictive models on AWS utilizing Docker containers for microservices architecture.",
                    "Lead a team of 3 junior analysts using Agile methodologies to deliver sprint goals."
                ]
            },
            {
                "title": "Data Analyst",
                "company": "InfoCorp Services",
                "duration": "2018 - 2021",
                "responsibilities": [
                    "Utilized Pandas and NumPy for extensive data cleaning and analysis of over 50 million records.",
                    "Designed SQL queries to optimize database retrieval times by 30%."
                ]
            }
        ],
        [{"degree": "Master of Science in Data Science", "institution": "Stanford University", "year": "2018"}],
        ["Python", "Machine Learning", "PyTorch", "Scikit-Learn", "Pandas", "NumPy", "SQL", "AWS", "Docker", "Git", "Agile", "Leadership"]
    )
    
    # 2. John Doe (Senior Python Developer - High Match) - DOCX
    create_docx_resume(
        "John_Doe_Python_Developer.docx",
        "John Doe",
        "john.doe@email.com",
        "+1-555-0102",
        "Senior Software Engineer specializing in backend development with Python. Enthusiastic about clean code, system architecture, database optimization, and cloud operations.",
        [
            {
                "title": "Lead Python Developer",
                "company": "SoftWorks Inc.",
                "duration": "2020 - Present",
                "responsibilities": [
                    "Architected and built microservices using Python, Django, FastAPI, and Flask.",
                    "Configured CI/CD pipelines using Git, Docker, and AWS ECS for automated testing and deployment.",
                    "Optimized database performance on PostgreSQL, improving system response times by 40%."
                ]
            },
            {
                "title": "Software Engineer",
                "company": "DevGroup Corp",
                "duration": "2016 - 2020",
                "responsibilities": [
                    "Developed REST APIs using Flask and integrated third-party payment gateways.",
                    "Collaborated with product teams in an Agile Scrum environment to deliver software increments."
                ]
            }
        ],
        [{"degree": "Bachelor of Science in Computer Science", "institution": "Massachusetts Institute of Technology", "year": "2016"}],
        ["Python", "Django", "FastAPI", "Flask", "SQL", "PostgreSQL", "AWS", "Docker", "Git", "CI/CD", "Agile", "Scrum", "REST API"]
    )

    # 3. Charlie Green (ML Engineer - High Match) - PDF
    create_pdf_resume(
        "Charlie_Green_ML_Engineer.pdf",
        "Charlie Green",
        "charlie.green@email.com",
        "+1-555-0103",
        "Machine Learning Engineer with strong theoretical foundations and practical skills in Deep Learning and Natural Language Processing. Experienced in training and deploying large-scale neural network models.",
        [
            {
                "title": "Machine Learning Engineer",
                "company": "AILabs",
                "duration": "2022 - Present",
                "responsibilities": [
                    "Developed NLP models for text categorization and sentiment analysis using PyTorch and HuggingFace.",
                    "Built ML pipelines using Scikit-Learn and managed deployments using Docker and Kubernetes.",
                    "Conducted research on new deep learning models to improve core system performance."
                ]
            },
            {
                "title": "Research Assistant",
                "company": "UC Berkeley AI Lab",
                "duration": "2021 - 2022",
                "responsibilities": [
                    "Implemented computer vision algorithms in Python and NumPy.",
                    "Collaborated on writing academic papers for machine learning conferences."
                ]
            }
        ],
        [{"degree": "Ph.D. in Computer Science (AI focus)", "institution": "University of California, Berkeley", "year": "2021"}],
        ["Python", "Machine Learning", "Deep Learning", "NLP", "PyTorch", "TensorFlow", "Scikit-Learn", "NumPy", "Docker", "Kubernetes", "Git"]
    )

    # 4. Alice Johnson (Backend Engineer - Medium Match) - TXT
    create_txt_resume(
        "Alice_Johnson_Backend_Engineer.txt",
        """ALICE JOHNSON
Email: alice.j@email.com | Phone: +1-555-0104 | Location: Austin, TX

PROFESSIONAL SUMMARY
Backend Engineer with 4 years of experience building reliable web applications. Passionate about Python, API development, and SQL database design. Excellent communicator and collaborative team player.

TECHNICAL SKILLS
Python, Flask, Django, SQL, MySQL, PostgreSQL, Git, REST APIs, HTML, CSS, Communication

WORK EXPERIENCE
Backend Developer | WebBuilders Studio | 2022 - Present
- Created robust REST APIs using Flask and Django to power e-commerce websites.
- Wrote complex SQL queries and optimized PostgreSQL databases.
- Collaborated in daily standups and Agile sprint planning.

Junior Developer | CodeBase Corp | 2020 - 2022
- Maintained backend Python scripts and fixed software bugs.
- Managed version control using Git and GitHub.

EDUCATION
Bachelor of Technology in Information Technology
Texas A&M University (2020)
"""
    )

    # 5. George Clark (Fullstack Engineer - Medium Match) - PDF
    create_pdf_resume(
        "George_Clark_Fullstack_Engineer.pdf",
        "George Clark",
        "george.clark@email.com",
        "+1-555-0105",
        "Fullstack Engineer with experience building responsive web applications from scratch. Capable in front-end frameworks like React and backend development using Python and Node.js.",
        [
            {
                "title": "Fullstack Developer",
                "company": "CloudApps Inc.",
                "duration": "2023 - Present",
                "responsibilities": [
                    "Developed web applications using React for frontend and Python (FastAPI) for backend services.",
                    "Utilized AWS services including S3 and RDS for cloud storage and deployment.",
                    "Managed source code using Git and performed unit testing."
                ]
            }
        ],
        [{"degree": "Bachelor of Engineering in Software Engineering", "institution": "University of Texas, Austin", "year": "2023"}],
        ["Python", "FastAPI", "React", "JavaScript", "Node.js", "SQL", "AWS", "Git", "HTML", "CSS"]
    )

    # 6. Diana Prince (DevOps Specialist - Medium Match) - TXT
    create_txt_resume(
        "Diana_Prince_DevOps_Specialist.txt",
        """DIANA PRINCE
Email: diana.p@email.com | Phone: +1-555-0106 | Location: Seattle, WA

PROFESSIONAL SUMMARY
DevOps Specialist with over 7 years of experience automating software deployments and managing cloud infrastructure. Strong focus on containerization, CI/CD pipelines, and high availability.

TECHNICAL SKILLS
Docker, Kubernetes, AWS, Jenkins, Git, Terraform, Ansible, Linux, Bash, CI/CD, Python

WORK EXPERIENCE
Senior DevOps Engineer | CloudOps Corp | 2021 - Present
- Automated cloud infrastructure provisioning using Terraform on AWS.
- Configured CI/CD pipelines using Jenkins and GitLab CI.
- Managed Kubernetes clusters and Docker containers in production.

Systems Administrator | SysNet Group | 2017 - 2021
- Maintained Linux servers and wrote automation scripts in Bash and Python.
- Monitored system performance and resolved outages.

EDUCATION
Bachelor of Science in Systems Engineering
University of Washington (2017)
"""
    )

    # 7. Evan Wright (Data Analyst - Low/Medium Match) - DOCX
    create_docx_resume(
        "Evan_Wright_Data_Analyst.docx",
        "Evan Wright",
        "evan.wright@email.com",
        "+1-555-0107",
        "Data Analyst with 2 years of experience turning complex datasets into visual dashboards. Strong skills in SQL, Python, and business intelligence tools. Effective communicator.",
        [
            {
                "title": "Junior Data Analyst",
                "company": "Metrics Analytics",
                "duration": "2024 - Present",
                "responsibilities": [
                    "Cleaned and manipulated large datasets using Python and Pandas.",
                    "Created interactive dashboards in Tableau and PowerBI.",
                    "Wrote SQL queries to extract data from data warehouses."
                ]
            }
        ],
        [{"degree": "Bachelor of Science in Statistics", "institution": "University of Illinois", "year": "2023"}],
        ["Python", "Pandas", "SQL", "Tableau", "Communication", "Time Management"]
    )

    # 8. Hannah Abbott (Java Developer - Low Match) - DOCX
    create_docx_resume(
        "Hannah_Abbott_Java_Developer.docx",
        "Hannah Abbott",
        "hannah.abbott@email.com",
        "+1-555-0108",
        "Java Developer with 4 years of experience building enterprise applications. Strong understanding of Object-Oriented Programming, MVC frameworks, and SQL databases.",
        [
            {
                "title": "Software Engineer",
                "company": "Enterprise Software Co.",
                "duration": "2022 - Present",
                "responsibilities": [
                    "Developed backend services using Java, Spring Boot, and Hibernate.",
                    "Designed database schemas and wrote stored procedures in MySQL.",
                    "Participated in agile ceremonies and code reviews."
                ]
            }
        ],
        [{"degree": "Bachelor of Science in Computer Science", "institution": "Purdue University", "year": "2022"}],
        ["Java", "Spring Boot", "Hibernate", "SQL", "MySQL", "Git", "OOP"]
    )

    # 9. Bob Brown (Frontend Developer - Low Match) - DOCX
    create_docx_resume(
        "Bob_Brown_Frontend_Developer.docx",
        "Bob Brown",
        "bob.brown@email.com",
        "+1-555-0109",
        "Frontend Developer specialized in creating beautiful and responsive user interfaces. Expert in HTML, CSS, JavaScript, React, and CSS frameworks like Tailwind CSS.",
        [
            {
                "title": "Frontend Developer",
                "company": "Creative Web Studio",
                "duration": "2023 - Present",
                "responsibilities": [
                    "Built responsive web pages using React and Tailwind CSS.",
                    "Collaborated with UI/UX designers to implement design prototypes.",
                    "Optimized web apps for maximum speed and mobile responsiveness."
                ]
            }
        ],
        [{"degree": "Bachelor of Arts in Graphic Design", "institution": "Rhode Island School of Design", "year": "2023"}],
        ["JavaScript", "React", "HTML", "CSS", "Tailwind", "Git", "Collaboration"]
    )

    # 10. Fiona Gallagher (Software Intern - Low Match) - TXT
    create_txt_resume(
        "Fiona_Gallagher_Software_Intern.txt",
        """FIONA GALLAGHER
Email: fiona.g@email.com | Phone: +1-555-0110 | Location: Chicago, IL

PROFESSIONAL SUMMARY
Motivated Computer Science undergraduate seeking a software engineering internship. Quick learner, passionate about Python programming, and enthusiastic about web technologies.

TECHNICAL SKILLS
Python, HTML, CSS, Git, Teamwork, Problem Solving

WORK EXPERIENCE
Software Intern | CodeCamp | 2025 - Present (Part-time)
- Developed simple script modules in Python to automate internal data file formatting.
- Contributed to styled landing pages using HTML and CSS.

EDUCATION
Undergraduate Student in Computer Science
University of Illinois Chicago (Expected Graduation: 2027)
"""
    )
    
    print("\nSuccessfully generated 10 diverse mock resumes in 'data/resumes/'.")

if __name__ == "__main__":
    main()
